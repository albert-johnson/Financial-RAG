from __future__ import annotations

import os
import time
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Optional

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

from src.config import (
    DATA_DIR,
    DEFAULT_CHUNK_SIZE,
    DEFAULT_CHUNK_OVERLAP,
)
from src.rag.vectorstore import VectorStoreManager
from src.rag.tabular_extractor import (
    extract_tables_from_pdf,
    extract_tables_from_docx,
    extract_tables_from_csv,
    extract_tables_from_xlsx,
    build_table_documents,
)
from src.rag.image_extractor import extract_pdf_page_images, extract_docx_inline_images


SUPPORTED_EXTS = {".pdf", ".docx", ".csv", ".xlsx"}


def _read_text_from_pdf(file_path: str) -> str:
    import pdfplumber
    texts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                try:
                    texts.append(page.extract_text() or "")
                except Exception:
                    continue
    except Exception:
        return ""
    return "\n".join(texts)


def _read_text_from_docx(file_path: str) -> str:
    from docx import Document as DocxDocument
    try:
        doc = DocxDocument(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception:
        return ""


def _read_text_from_csv(file_path: str) -> str:
    import csv
    try:
        with open(file_path, newline="") as f:
            reader = csv.reader(f)
            rows = list(reader)
        return "\n".join([",".join(row) for row in rows])
    except Exception:
        return ""


def _read_text_from_xlsx(file_path: str) -> str:
    try:
        import openpyxl
    except Exception:
        return ""
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        texts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append(["" if v is None else str(v) for v in row])
            texts.append("Sheet: " + sheet + "\n" + "\n".join([",".join(r) for r in rows]))
        return "\n\n".join(texts)
    except Exception:
        return ""


def _split_into_documents(text: str, source_path: str) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=DEFAULT_CHUNK_SIZE,
        chunk_overlap=DEFAULT_CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", ".", " "]
    )
    chunks = splitter.split_text(text)
    documents: List[Document] = []
    for i, chunk in enumerate(chunks):
        metadata = {
            "source": os.path.abspath(source_path),
            "doc_chunk_id": f"{os.path.abspath(source_path)}::chunk::{i}",
            "chunk_index": i,
        }
        documents.append(Document(page_content=chunk, metadata=metadata))
    return documents


def ingest_file(path: str, vsm: Optional[VectorStoreManager] = None) -> Dict:
    vsm = vsm or VectorStoreManager()
    ext = Path(path).suffix.lower()
    if ext not in SUPPORTED_EXTS:
        return {"status": "skipped", "reason": f"unsupported extension {ext}", "path": path}

    abs_path = os.path.abspath(path)

    # Remove old entries for this source
    vsm.delete_documents_by_source(abs_path)
    vsm.delete_tables_by_source(abs_path)

    # Extract text
    if ext == ".pdf":
        text = _read_text_from_pdf(abs_path)
    elif ext == ".docx":
        text = _read_text_from_docx(abs_path)
    elif ext == ".csv":
        text = _read_text_from_csv(abs_path)
    elif ext == ".xlsx":
        text = _read_text_from_xlsx(abs_path)
    else:
        text = ""

    # Chunk and add to vector store
    docs = _split_into_documents(text, abs_path) if text else []
    if docs:
        vsm.add_documents(docs)

    # Extract tables and add their descriptions to table index
    extracted_tables = []
    if ext == ".pdf":
        extracted_tables = extract_tables_from_pdf(abs_path)
    elif ext == ".docx":
        extracted_tables = extract_tables_from_docx(abs_path)
    elif ext == ".csv":
        extracted_tables = extract_tables_from_csv(abs_path)
    elif ext == ".xlsx":
        extracted_tables = extract_tables_from_xlsx(abs_path)

    table_docs = build_table_documents(extracted_tables)
    if table_docs:
        vsm.add_table_descriptions(table_docs)

    # Extract page images / snapshots
    snapshots: List[str] = []
    if ext == ".pdf":
        snapshots = extract_pdf_page_images(abs_path)
    elif ext == ".docx":
        snapshots = extract_docx_inline_images(abs_path)

    vsm.persist()

    return {
        "status": "ingested",
        "path": abs_path,
        "num_chunks": len(docs),
        "num_tables": len(extracted_tables),
        "snapshots": snapshots,
    }


def ingest_paths(paths: List[str]) -> List[Dict]:
    vsm = VectorStoreManager()
    results = []
    for p in paths:
        results.append(ingest_file(p, vsm=vsm))
    return results


def ingest_data_dir() -> List[Dict]:
    paths: List[str] = []
    for root, _, files in os.walk(DATA_DIR):
        for f in files:
            if Path(f).suffix.lower() in SUPPORTED_EXTS:
                paths.append(os.path.join(root, f))
    return ingest_paths(paths)