from __future__ import annotations

import os
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple, Dict

import csv
import pdfplumber
from docx import Document as DocxDocument
from langchain.schema import Document

from src.config import TABLES_DIR

try:
    import pyarrow as pa
    import pyarrow.parquet as pq
except Exception:  # fallback type stubs
    pa = None
    pq = None

try:
    import openpyxl
except Exception:
    openpyxl = None


@dataclass
class ExtractedTable:
    table_id: str
    source: str
    location: str  # e.g., "page 3", "sheet:Sheet1", "table 1"
    path: str      # path to saved csv
    num_rows: int
    num_cols: int
    columns: List[str]


def _save_rows(headers: List[str], rows: List[List[str]], base_name: str) -> str:
    TABLES_DIR.mkdir(parents=True, exist_ok=True)
    table_path = TABLES_DIR / f"{base_name}.csv"
    with open(table_path, "w", newline="") as f:
        writer = csv.writer(f)
        if headers:
            writer.writerow(headers)
        writer.writerows(rows)
    return str(table_path)


def extract_tables_from_pdf(file_path: str) -> List[ExtractedTable]:
    tables: List[ExtractedTable] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    page_tables = page.extract_tables() or []
                except Exception:
                    page_tables = []
                for t_index, table in enumerate(page_tables, start=1):
                    try:
                        # table is list of rows (list[str|None])
                        if not table:
                            continue
                        # Determine headers
                        headers = [str(x) if x is not None else "" for x in table[0]]
                        data_rows = [
                            [str(x) if x is not None else "" for x in row]
                            for row in table[1:]
                        ]
                        table_uuid = str(uuid.uuid4())
                        base_name = f"pdf_{Path(file_path).stem}_p{page_num}_t{t_index}_{table_uuid[:8]}"
                        path = _save_rows(headers, data_rows, base_name)
                        num_cols = len(headers)
                        tables.append(
                            ExtractedTable(
                                table_id=table_uuid,
                                source=os.path.abspath(file_path),
                                location=f"page {page_num}",
                                path=path,
                                num_rows=len(data_rows),
                                num_cols=num_cols,
                                columns=headers,
                            )
                        )
                    except Exception:
                        continue
    except Exception:
        pass
    return tables


def extract_tables_from_docx(file_path: str) -> List[ExtractedTable]:
    tables: List[ExtractedTable] = []
    try:
        doc = DocxDocument(file_path)
        for t_index, table in enumerate(doc.tables, start=1):
            data = [[cell.text for cell in row.cells] for row in table.rows]
            if not data:
                continue
            headers = [str(x) for x in data[0]]
            data_rows = [[str(x) for x in row] for row in data[1:]]
            table_uuid = str(uuid.uuid4())
            base_name = f"docx_{Path(file_path).stem}_t{t_index}_{table_uuid[:8]}"
            path = _save_rows(headers, data_rows, base_name)
            tables.append(
                ExtractedTable(
                    table_id=table_uuid,
                    source=os.path.abspath(file_path),
                    location=f"table {t_index}",
                    path=path,
                    num_rows=len(data_rows),
                    num_cols=len(headers),
                    columns=headers,
                )
            )
    except Exception:
        pass
    return tables


def extract_tables_from_csv(file_path: str) -> List[ExtractedTable]:
    try:
        with open(file_path, newline="") as f:
            reader = csv.reader(f)
            rows = list(reader)
    except Exception:
        return []
    headers = [str(x) for x in rows[0]] if rows else []
    data_rows = rows[1:] if len(rows) > 1 else []
    table_uuid = str(uuid.uuid4())
    base_name = f"csv_{Path(file_path).stem}_{table_uuid[:8]}"
    path = _save_rows(headers, data_rows, base_name)
    return [
        ExtractedTable(
            table_id=table_uuid,
            source=os.path.abspath(file_path),
            location="file",
            path=path,
            num_rows=len(data_rows),
            num_cols=len(headers),
            columns=headers,
        )
    ]


def extract_tables_from_xlsx(file_path: str) -> List[ExtractedTable]:
    tables: List[ExtractedTable] = []
    if openpyxl is None:
        return tables
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append(["" if v is None else str(v) for v in row])
            if not rows:
                continue
            headers = [str(x) for x in rows[0]]
            data_rows = rows[1:] if len(rows) > 1 else []
            table_uuid = str(uuid.uuid4())
            base_name = f"xlsx_{Path(file_path).stem}_{sheet_name}_{table_uuid[:8]}"
            path = _save_rows(headers, data_rows, base_name)
            tables.append(
                ExtractedTable(
                    table_id=table_uuid,
                    source=os.path.abspath(file_path),
                    location=f"sheet:{sheet_name}",
                    path=path,
                    num_rows=len(data_rows),
                    num_cols=len(headers),
                    columns=headers,
                )
            )
    except Exception:
        pass
    return tables


def build_table_documents(extracted: List[ExtractedTable]) -> List[Document]:
    docs: List[Document] = []
    for t in extracted:
        description = (
            f"Table from {os.path.basename(t.source)} at {t.location}. "
            f"Rows: {t.num_rows}, Cols: {t.num_cols}. Columns: {', '.join(t.columns[:12])}"
        )
        metadata = {
            "source": t.source,
            "table_id": t.table_id,
            "location": t.location,
            "table_path": t.path,
            "columns": t.columns,
            "num_rows": t.num_rows,
            "num_cols": t.num_cols,
        }
        docs.append(Document(page_content=description, metadata=metadata))
    return docs


def load_table_df(table_path: str):
    # Returns a simple dict structure usable for plotting
    ext = Path(table_path).suffix.lower()
    headers: List[str] = []
    data_rows: List[List[str]] = []
    if ext == ".csv":
        with open(table_path, newline="") as f:
            reader = csv.reader(f)
            rows = list(reader)
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []
        return {"columns": headers, "rows": data_rows}
    # Fallback empty
    return {"columns": [], "rows": []}